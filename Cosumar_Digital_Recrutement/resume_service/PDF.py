import easyocr
import re
from pdf2image import convert_from_bytes
import numpy as np
import io
import fitz 
from typing import Dict, Optional
from docx import Document
from docx2pdf import convert
import tempfile
import os
import zipfile
import xml.etree.ElementTree as ET

def extract_emails(text):
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    return re.findall(email_pattern, text)

def extract_phones(text):
    phone_patterns = [
        r'(\+212|0)[-\s]*([67])[-\s]*(\d{2})[-\s]*(\d{2})[-\s]*(\d{2})[-\s]*(\d{2})',
        r'(\+212|0)([67]\d{8})',
        r'(\d{2})[-\s]*(\d{2})[-\s]*(\d{2})[-\s]*(\d{2})[-\s]*(\d{2})',
        r'\+212[-\s]*\d[-\s]*\d{2}[-\s]*\d{3}[-\s]*\d{3}',
        r'0[67][-\s]*\d{2}[-\s]*\d{2}[-\s]*\d{2}[-\s]*\d{2}'
    ]
    
    phones = []
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple) and len(match) > 1:
                if len(match) == 6:
                    phone = f"{match[0]}{match[1]}{match[2]}{match[3]}{match[4]}{match[5]}"
                elif len(match) == 2:
                    phone = f"{match[0]}{match[1]}"
                elif len(match) == 5:
                    phone = '0' + ''.join(match)
                else:
                    phone = ''.join(match)
            else:
                phone = str(match).replace('-', '').replace(' ', '')
                
            if phone and len(phone) >= 9:
                phones.append(phone)
    
    direct_patterns = [
        r'\+212[-\s]*[67][-\s]*\d{2}[-\s]*\d{3}[-\s]*\d{3}',
        r'\+212[-\s]*[67][-\s]*\d{2}[-\s]*\d{2}[-\s]*\d{2}[-\s]*\d{2}',
        r'0[67][-\s]*\d{2}[-\s]*\d{2}[-\s]*\d{2}[-\s]*\d{2}'
    ]
    
    for pattern in direct_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            clean_phone = match.replace('-', '').replace(' ', '')
            if clean_phone and len(clean_phone) >= 9:
                phones.append(clean_phone)
    
    return list(set(phones))

def extract_cv_data(pdf_bytes, lang='fr'):
    try:
        poppler_path = r"C:/Program Files/poppler-24.08.0/Library/bin"
        
        if isinstance(pdf_bytes, bytes):
            images = convert_from_bytes(pdf_bytes, dpi=300, poppler_path=poppler_path)
        else:
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_bytes, dpi=300, poppler_path=poppler_path)
        
        reader = easyocr.Reader([lang])
        
        email = None
        phone = None
        all_text = ""
        
        for i, image in enumerate(images):
            img_array = np.array(image)
            
            result = reader.readtext(img_array, detail=0, paragraph=False)
            
            page_text = " ".join(result)
            all_text += page_text + " "
            
            if email is None:
                page_emails = extract_emails(page_text)
                if page_emails:
                    email = page_emails[0]
            
            if phone is None:
                page_phones = extract_phones(page_text)
                if page_phones:
                    phone = page_phones[0]
            
            if email and phone:
                break
        
        if email is None:
            all_emails = extract_emails(all_text)
            if all_emails:
                email = all_emails[0]
        
        if phone is None:
            all_phones = extract_phones(all_text)
            if all_phones:
                phone = all_phones[0]
        
        return {
            'email': email,
            'phone': phone,
        }
        
    except Exception as e:
        return {
            'email': None,
            'phone': None,
        }
    
def replace_word_in_pdf(pdf_bytes: bytes, old_word: str, new_word: str) -> bytes:
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        page = pdf_document[0]
        
        # Search for the text you want to replace
        areas = page.search_for(old_word)
        
        for area in areas:
            # Calculate appropriate font size based on text length and available space
            area_width = area.width
            area_height = area.height
            
            # Start with a reasonable font size
            font_size = 10
            
            # Adjust font size based on text length and available width
            if len(new_word) > 15:
                font_size = 8
            elif len(new_word) > 25:
                font_size = 7
            elif len(new_word) > 35:
                font_size = 6
            
            # Make sure font size fits in the height
            if area_height < font_size:
                font_size = max(6, int(area_height * 0.8))
            
            # First, redact (cover with white rectangle) - make it slightly larger
            expanded_area = fitz.Rect(
                area.x0 - 2, 
                area.y0 - 2, 
                area.x1 + len(new_word) * 2,  # Expand width based on text length
                area.y1 + 2
            )
            page.add_redact_annot(expanded_area, fill=(1, 1, 1))
            page.apply_redactions()
            
            # Then insert the new text at the same spot
            page.insert_text(
                area.bl, 
                new_word, 
                fontsize=font_size, 
                color=(0, 0, 0)
            )
    
        output_buffer = io.BytesIO()
        pdf_document.save(output_buffer)
        pdf_document.close()
        
        output_buffer.seek(0)
        return output_buffer.getvalue()
        
    except Exception as e:
        print(f"Error replacing word in PDF: {str(e)}")
        return pdf_bytes  # Return original if error occurs


def replace_multiple_words_in_pdf(pdf_bytes: bytes, replacements: Dict[str, str]) -> bytes:
    modified_pdf = pdf_bytes
    
    for old_word, new_word in replacements.items():
        modified_pdf = replace_word_in_pdf(modified_pdf, old_word, new_word)
    
    return modified_pdf

def create_docx_from_template(docx_path: str, replacements: Dict[str, str], output_path: str = None) -> bytes:
    """
    Create a DOCX from a DOCX template by replacing placeholders while preserving all formatting
    
    Args:
        docx_path: Path to the DOCX template file
        replacements: Dictionary of {placeholder: replacement_value}
        output_path: Optional path to save the DOCX file
    
    Returns:
        DOCX content as bytes
    """
    try:
        # Load the DOCX document
        doc = Document(docx_path)
        
        # Function to safely replace text in individual runs without destroying formatting
        def safe_replace_in_paragraph(paragraph):
            # Build a list of all text content and their corresponding runs
            run_texts = []
            for run in paragraph.runs:
                run_texts.append(run.text)
            
            # Combine all text to search for placeholders
            full_text = ''.join(run_texts)
            original_full_text = full_text
            
            # Apply all replacements to the combined text
            for placeholder, replacement in replacements.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, replacement)
                    print(f"✅ Found and replacing {placeholder} with {replacement}")
            
            # Only modify if text actually changed
            if full_text != original_full_text:
                # Find which runs need to be updated
                # Simple approach: put all the new text in the first run and clear others
                if paragraph.runs:
                    # Keep the first run's formatting and put all text there
                    paragraph.runs[0].text = full_text
                    
                    # Clear the text from other runs but keep their formatting intact
                    for i in range(1, len(paragraph.runs)):
                        paragraph.runs[i].text = ""
                        
        # Process paragraphs - but only those that contain our placeholders
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if any(placeholder in paragraph_text for placeholder in replacements.keys()):
                safe_replace_in_paragraph(paragraph)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = paragraph.text
                        if any(placeholder in paragraph_text for placeholder in replacements.keys()):
                            safe_replace_in_paragraph(paragraph)
        
        # Process headers and footers (more carefully)
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    paragraph_text = paragraph.text
                    if any(placeholder in paragraph_text for placeholder in replacements.keys()):
                        safe_replace_in_paragraph(paragraph)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    paragraph_text = paragraph.text
                    if any(placeholder in paragraph_text for placeholder in replacements.keys()):
                        safe_replace_in_paragraph(paragraph)
        
        # Save the document
        if output_path:
            # Save directly to the specified output path
            doc.save(output_path)
            
            # Read the DOCX content
            with open(output_path, 'rb') as docx_file:
                docx_bytes = docx_file.read()
        else:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                doc.save(temp_docx.name)
                
                # Read the DOCX content
                with open(temp_docx.name, 'rb') as docx_file:
                    docx_bytes = docx_file.read()
                
                # Clean up temporary file
                os.unlink(temp_docx.name)
        
        print(f"🎉 Successfully created DOCX from template")
        return docx_bytes
        
    except Exception as e:
        print(f"❌ Error creating DOCX from template: {str(e)}")
        import traceback
        traceback.print_exc()
        return b''  # Return empty bytes if error occurs

def create_pdf_from_docx_template_xml(docx_path: str, replacements: Dict[str, str], output_path: str = None) -> bytes:
    """
    Create a PDF from DOCX template by working at XML level to preserve ALL document elements, then convert to PDF
    
    Args:
        docx_path: Path to the DOCX template file
        replacements: Dictionary of {placeholder: replacement_value}
        output_path: Optional path to save the PDF file
    
    Returns:
        PDF content as bytes
    """
    try:
        # Copy the original file to a temporary location
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
            
        # Copy original file
        import shutil
        shutil.copy2(docx_path, temp_docx_path)
        
        # Open as ZIP (DOCX is a ZIP file)
        with zipfile.ZipFile(temp_docx_path, 'r') as zip_read:
            # Extract all files
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_read.extractall(temp_dir)
                
                # Process the main document XML
                document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
                if os.path.exists(document_xml_path):
                    with open(document_xml_path, 'r', encoding='utf-8') as f:
                        xml_content = f.read()
                    
                    # Replace placeholders in XML content
                    for placeholder, replacement in replacements.items():
                        if placeholder in xml_content:
                            xml_content = xml_content.replace(placeholder, replacement)
                            print(f"✅ Replaced {placeholder} with {replacement} in document XML")
                    
                    # Write back the modified XML
                    with open(document_xml_path, 'w', encoding='utf-8') as f:
                        f.write(xml_content)
                
                # Process headers if they exist
                headers_dir = os.path.join(temp_dir, 'word')
                for file_name in os.listdir(headers_dir):
                    if file_name.startswith('header') and file_name.endswith('.xml'):
                        header_path = os.path.join(headers_dir, file_name)
                        with open(header_path, 'r', encoding='utf-8') as f:
                            xml_content = f.read()
                        
                        for placeholder, replacement in replacements.items():
                            if placeholder in xml_content:
                                xml_content = xml_content.replace(placeholder, replacement)
                                print(f"✅ Replaced {placeholder} with {replacement} in {file_name}")
                        
                        with open(header_path, 'w', encoding='utf-8') as f:
                            f.write(xml_content)
                
                # Process footers if they exist
                for file_name in os.listdir(headers_dir):
                    if file_name.startswith('footer') and file_name.endswith('.xml'):
                        footer_path = os.path.join(headers_dir, file_name)
                        with open(footer_path, 'r', encoding='utf-8') as f:
                            xml_content = f.read()
                        
                        for placeholder, replacement in replacements.items():
                            if placeholder in xml_content:
                                xml_content = xml_content.replace(placeholder, replacement)
                                print(f"✅ Replaced {placeholder} with {replacement} in {file_name}")
                        
                        with open(footer_path, 'w', encoding='utf-8') as f:
                            f.write(xml_content)
                
                # Create new DOCX file with modified content
                filled_docx_path = temp_docx_path + '_filled.docx'
                with zipfile.ZipFile(filled_docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arc_name = os.path.relpath(file_path, temp_dir)
                            zip_write.write(file_path, arc_name)
        
        # Convert the filled DOCX to PDF
        if output_path:
            # Convert directly to the specified output path
            convert(filled_docx_path, output_path)
            pdf_output_path = output_path
        else:
            # Convert to temporary PDF file
            pdf_output_path = temp_docx_path + '_output.pdf'
            convert(filled_docx_path, pdf_output_path)
        
        # Read the PDF content
        with open(pdf_output_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Clean up temporary files
        os.unlink(temp_docx_path)
        os.unlink(filled_docx_path)
        if not output_path:
            os.unlink(pdf_output_path)
        
        print(f"🎉 Successfully created PDF from DOCX template using XML method")
        return pdf_bytes
        
    except Exception as e:
        print(f"❌ Error creating PDF from DOCX using XML method: {str(e)}")
        import traceback
        traceback.print_exc()
        return b''

def create_docx_from_template_xml(docx_path: str, replacements: Dict[str, str]) -> bytes:
    """
    Create a filled DOCX from DOCX template by working at XML level to preserve ALL document elements
    
    Args:
        docx_path: Path to the DOCX template file
        replacements: Dictionary of {placeholder: replacement_value}
    
    Returns:
        DOCX content as bytes
    """
    try:
        # Copy the original file to a temporary location
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
            
        # Copy original file
        import shutil
        shutil.copy2(docx_path, temp_docx_path)
        
        # Open as ZIP (DOCX is a ZIP file)
        with zipfile.ZipFile(temp_docx_path, 'r') as zip_read:
            # Extract all files
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_read.extractall(temp_dir)
                
                # Process the main document XML
                document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
                if os.path.exists(document_xml_path):
                    with open(document_xml_path, 'r', encoding='utf-8') as f:
                        xml_content = f.read()
                    
                    # Debug: Find all placeholders in the template
                    import re
                    found_placeholders = re.findall(r'«[^»]+»', xml_content)
                    print(f"🔍 Debug: Found placeholders in template: {found_placeholders}")
                    
                    # Also check for potential encoding variations
                    alt_placeholders = re.findall(r'<<[^>]+>>', xml_content)  # Alternative format
                    if alt_placeholders:
                        print(f"🔍 Debug: Found alternative placeholders: {alt_placeholders}")
                    
                    # Check for SIGNATURE_ENCADRANT specifically in different formats
                    signature_variations = [
                        '«SIGNATURE_ENCADRANT»',
                        '<<SIGNATURE_ENCADRANT>>',
                        'SIGNATURE_ENCADRANT',
                        'signature_encadrant',
                        'Signature_Encadrant'
                    ]
                    
                    print(f"🔍 Debug: Checking for SIGNATURE_ENCADRANT variations:")
                    for variation in signature_variations:
                        count = xml_content.count(variation)
                        if count > 0:
                            print(f"   ✅ Found '{variation}': {count} times")
                        else:
                            print(f"   ❌ Not found: '{variation}'")
                    
                    # Show a sample of the XML content to see the structure
                    xml_lines = xml_content.split('\n')
                    print(f"🔍 Debug: XML content sample (first 10 lines):")
                    for i, line in enumerate(xml_lines[:10]):
                        print(f"   {i+1}: {line[:100]}...")
                        
                    # Look for any text that contains "SIGNATURE" or "ENCADRANT"
                    signature_matches = re.findall(r'[^>]*(?:SIGNATURE|ENCADRANT)[^<]*', xml_content, re.IGNORECASE)
                    if signature_matches:
                        print(f"🔍 Debug: Found text containing SIGNATURE/ENCADRANT:")
                        for match in signature_matches[:5]:  # Show first 5 matches
                            print(f"   '{match}'")
                    
                    # Replace placeholders in XML content with special formatting
                    replaced_count = 0
                    
                    # First, try to consolidate split placeholders in the XML
                    # This handles cases where placeholders are split across multiple <w:t> elements
                    print(f"🔄 Debug: Attempting to consolidate split placeholders...")
                    
                    # Look for patterns where placeholders might be split
                    # Pattern: text ending with part of placeholder + closing tag + opening tag + rest of placeholder
                    split_patterns = [
                        (r'(«[^»]*)</w:t>([^<]*<w:t[^>]*>)([^<]*»)', r'\1\3'),  # «PART</w:t>...<w:t>REST»
                        (r'(«[A-Z_]*)</w:t>([^<]*<w:t[^>]*>)([A-Z_]*»)', r'\1\3'),  # More specific pattern
                    ]
                    
                    for pattern, replacement in split_patterns:
                        matches = re.findall(pattern, xml_content)
                        if matches:
                            print(f"🔍 Debug: Found {len(matches)} split placeholder patterns")
                            xml_content = re.sub(pattern, replacement, xml_content)
                    
                    # Check again after consolidation
                    found_placeholders_after = re.findall(r'«[^»]+»', xml_content)
                    print(f"🔍 Debug: Placeholders after consolidation: {found_placeholders_after}")
                    
                    for placeholder, replacement in replacements.items():
                        original_count = xml_content.count(placeholder)
                        if original_count > 0:
                            # Special handling for SIGNATURE_ENCADRANT to make it bold
                            if placeholder == '«SIGNATURE_ENCADRANT»':
                                # Use regex to find and replace within w:t tags with bold formatting
                                pattern = r'(<w:t[^>]*>)([^<]*' + re.escape(placeholder) + r'[^<]*)(<\/w:t>)'
                                
                                def make_bold_replacement(match):
                                    opening_tag = match.group(1)  # <w:t> or <w:t xml:space="preserve">
                                    text_content = match.group(2)  # Text containing placeholder
                                    closing_tag = match.group(3)  # </w:t>
                                    
                                    # Split the text around the placeholder
                                    before_placeholder = text_content.split(placeholder)[0]
                                    after_placeholder = ''.join(text_content.split(placeholder)[1:])
                                    
                                    # Create the replacement with bold formatting
                                    # Structure: existing_text</w:t></w:r><w:r><w:rPr><w:b/></w:rPr><w:t>BOLD_TEXT</w:t></w:r><w:r><w:t>remaining_text
                                    return f'{opening_tag}{before_placeholder}</w:t></w:r><w:r><w:rPr><w:b/></w:rPr><w:t>{replacement}</w:t></w:r><w:r><w:t>{after_placeholder}</w:t>'
                                
                                xml_content = re.sub(pattern, make_bold_replacement, xml_content)
                                print(f"✅ Replaced {original_count} instances of {placeholder} with BOLD '{replacement}' in document XML")
                            else:
                                xml_content = xml_content.replace(placeholder, replacement)
                                print(f"✅ Replaced {original_count} instances of {placeholder} with '{replacement}' in document XML")
                            replaced_count += original_count
                        else:
                            print(f"⚠️ Placeholder {placeholder} not found in document XML")
                    
                    print(f"📊 Debug: Total replacements made: {replaced_count}")
                    
                    # Final check: see what placeholders remain
                    remaining_placeholders = re.findall(r'«[^»]+»', xml_content)
                    if remaining_placeholders:
                        print(f"⚠️ Remaining unreplaced placeholders: {remaining_placeholders}")
                    
                    # Write back the modified XML
                    with open(document_xml_path, 'w', encoding='utf-8') as f:
                        f.write(xml_content)
                
                # Process headers if they exist
                headers_dir = os.path.join(temp_dir, 'word')
                for file_name in os.listdir(headers_dir):
                    if file_name.startswith('header') and file_name.endswith('.xml'):
                        header_path = os.path.join(headers_dir, file_name)
                        with open(header_path, 'r', encoding='utf-8') as f:
                            xml_content = f.read()
                        
                        for placeholder, replacement in replacements.items():
                            if placeholder in xml_content:
                                xml_content = xml_content.replace(placeholder, replacement)
                                print(f"✅ Replaced {placeholder} with {replacement} in {file_name}")
                        
                        with open(header_path, 'w', encoding='utf-8') as f:
                            f.write(xml_content)
                
                # Process footers if they exist
                for file_name in os.listdir(headers_dir):
                    if file_name.startswith('footer') and file_name.endswith('.xml'):
                        footer_path = os.path.join(headers_dir, file_name)
                        with open(footer_path, 'r', encoding='utf-8') as f:
                            xml_content = f.read()
                        
                        for placeholder, replacement in replacements.items():
                            if placeholder in xml_content:
                                xml_content = xml_content.replace(placeholder, replacement)
                                print(f"✅ Replaced {placeholder} with {replacement} in {file_name}")
                        
                        with open(footer_path, 'w', encoding='utf-8') as f:
                            f.write(xml_content)
                
                # Create new DOCX file with modified content
                filled_docx_path = temp_docx_path + '_filled.docx'
                with zipfile.ZipFile(filled_docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arc_name = os.path.relpath(file_path, temp_dir)
                            zip_write.write(file_path, arc_name)
        
        # Read the filled DOCX content
        with open(filled_docx_path, 'rb') as f:
            docx_bytes = f.read()
        
        # Clean up temporary files
        os.unlink(temp_docx_path)
        os.unlink(filled_docx_path)
        
        print(f"🎉 Successfully created DOCX from template using XML method")
        return docx_bytes
        
    except Exception as e:
        print(f"❌ Error creating DOCX from template using XML method: {str(e)}")
        import traceback
        traceback.print_exc()
        return b''

def convert_docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX bytes to PDF bytes
    
    Args:
        docx_bytes: DOCX content as bytes
    
    Returns:
        PDF content as bytes
    """
    try:
        # Initialize COM for the current thread (Windows requirement)
        import pythoncom
        pythoncom.CoInitialize()
        
        try:
            # Create temporary files
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx_file:
                temp_docx_path = temp_docx_file.name
                temp_docx_file.write(docx_bytes)
            
            # Convert DOCX to PDF
            temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
            convert(temp_docx_path, temp_pdf_path)
            
            # Read PDF content
            with open(temp_pdf_path, 'rb') as f:
                pdf_bytes = f.read()
            
            # Clean up temporary files
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
            
            print(f"🎉 Successfully converted DOCX bytes to PDF bytes")
            return pdf_bytes
            
        finally:
            # Always uninitialize COM
            pythoncom.CoUninitialize()
            
    except Exception as e:
        print(f"❌ Error converting DOCX bytes to PDF bytes: {str(e)}")
        import traceback
        traceback.print_exc()
        return b''

if __name__ == "__main__":
    pdf_path = "C:/Users/othma/Downloads/CV_FR-9.pdf"
    
    with open(pdf_path, 'rb') as f:
        pdf_bytes = f.read()
    
    data = extract_cv_data(pdf_bytes)
    print(f"🎉 Extraction complete:")
    print(data)